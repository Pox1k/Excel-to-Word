import tkinter as tk
from tkinter import filedialog, messagebox, Canvas, Scrollbar, Frame
from openpyxl import load_workbook
from docx import Document

def select_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
    excel_entry.delete(0, tk.END)
    excel_entry.insert(0, file_path)

def select_word_template():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    word_entry.delete(0, tk.END)
    word_entry.insert(0, file_path)

def select_output_folder():
    folder_path = filedialog.askdirectory()
    output_entry.delete(0, tk.END)
    output_entry.insert(0, folder_path)

def generate_document():
    excel_path = excel_entry.get()
    word_template_path = word_entry.get()
    output_folder = output_entry.get()
    
    if not excel_path or not word_template_path or not output_folder:
        messagebox.showwarning("Missing information", "Please provide all file paths.")
        return
    
    try:
        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active
        base_value = ws['L43'].value
    except Exception as e:
        messagebox.showerror("Error", f"Error reading Excel file: {e}")
        return
    
    if base_value is None:
        messagebox.showerror("Error", "The value in cell L43 is empty or invalid.")
        return

    value_200_plus = base_value + 200
    value_400_plus = base_value + 400
    value_200_minus = base_value - 200
    value_400_minus = base_value - 400

    data_dict = {
        'word1': base_value,
        'word2': value_200_plus,
        'word3': value_400_plus,
        'word4': value_200_minus,
        'word5': value_400_minus,
    }

    try:
        doc = Document(word_template_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error loading Word template: {e}")
        return

    for name, var in presence_vars.items():
        key = presence_keys[name]
        replacement = "+" if var.get() else "-"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, replacement)
    
    for name, time_entry in time_entries.items():
        key = time_keys[name]
        time_value = time_entry.get() if time_entry.get() else "-"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, time_value)

    def replace_values(data_dict):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data_dict.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

    replace_values(data_dict)

    output_file_path = filedialog.asksaveasfilename(defaultextension=".docx", initialdir=output_folder, filetypes=[("Word files", "*.docx")])
    if not output_file_path:
        messagebox.showwarning("No file selected", "Please provide a file name for the output.")
        return

    try:
        doc.save(output_file_path)
        messagebox.showinfo("Success", "Document generated successfully!")
    except Exception as e:
        messagebox.showerror("Error", f"Error saving document: {e}")

root = tk.Tk()
root.title("Excel to Word Automation")

tk.Label(root, text="Select Excel File:").grid(row=0, column=0, padx=10, pady=5)
excel_entry = tk.Entry(root, width=50)
excel_entry.grid(row=0, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=select_excel_file).grid(row=0, column=2, padx=10, pady=5)

tk.Label(root, text="Select Word Template:").grid(row=1, column=0, padx=10, pady=5)
word_entry = tk.Entry(root, width=50)
word_entry.grid(row=1, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=select_word_template).grid(row=1, column=2, padx=10, pady=5)

tk.Label(root, text="Select Output Folder:").grid(row=2, column=0, padx=10, pady=5)
output_entry = tk.Entry(root, width=50)
output_entry.grid(row=2, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=select_output_folder).grid(row=2, column=2, padx=10, pady=5)

tk.Button(root, text="Generate Document", command=generate_document).grid(row=3, column=1, padx=10, pady=20)

canvas = Canvas(root)
scrollbar = Scrollbar(root, orient="vertical", command=canvas.yview)
scrollable_frame = Frame(canvas)

scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(
        scrollregion=canvas.bbox("all")
    )
)

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.grid(row=4, column=0, columnspan=3, padx=10, pady=5)
scrollbar.grid(row=4, column=3, sticky="ns")

presence_vars = {}
presence_keys = {}
time_entries = {}
time_keys = {}
names = [
    " Сергій", " Михайло", " Микола", " Михайло", " Андрій",
    " Микола", " Віталій", " Станіслав", " Олександр", " Віктор",
    " Олег", " Наталія", " Ігор", " Євген", " Станіслав",
    " Андрій", " Олександр", " Олександр", " Людмила", " Василь",
    " Ярослав", " Андрій", " Сергій", " Людмила", " Олександр", 
    " Сергій", " Юрій", " Роман", " Юрій", " Віктор", 
    " Андрій", " Сергій", " Руслан", " Володимир", " Іван", 
    " Володимир", " Дмитро", " Віталій", " Наталія", " Валентин", 
    " Василь"
]

for index, name in enumerate(names):
    var = tk.IntVar()
    chk = tk.Checkbutton(scrollable_frame, text=name, variable=var)
    chk.grid(row=index, column=0, padx=10, pady=2, sticky="w")
    presence_vars[name] = var
    
    if index < 9:
        presence_keys[name] = f"work{index + 1}"
        time_key = f"time{index + 1}"
    elif index < 18:
        presence_keys[name] = f"list{index - 8}"
        time_key = f"nide{index - 8}"
    elif index < 27:
        presence_keys[name] = f"scan{index - 17}"
        time_key = f"extra{index - 17}"
    elif index < 36:
        presence_keys[name] = f"check{index - 26}"
        time_key = f"temp{index - 26}"
    elif index < 41:
        presence_keys[name] = f"mark{index - 35}"
        time_key = f"tick{index - 35}"
    elif index < 43:
        presence_keys[name] = f"mark{index - 34}"
        time_key = f"tick{index - 34}"
    else:
        presence_keys[name] = "unknown"
        time_key = "unknown"
    
    time_entry = tk.Entry(scrollable_frame, width=10)
    time_entry.grid(row=index, column=1, padx=10, pady=2)
    time_entries[name] = time_entry
    time_keys[name] = time_key

root.mainloop()
